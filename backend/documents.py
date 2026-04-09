"""
Extract text from .txt / .docx / .pdf and build downloadable .txt / .docx / .pdf outputs.
"""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path
from typing import Literal

import fpdf
from docx import Document
from fpdf import FPDF
from pypdf import PdfReader

MAX_INPUT_CHARS = 50_000
MAX_UPLOAD_BYTES = 8 * 1024 * 1024

OutputFormat = Literal["txt", "docx", "pdf"]


def safe_stem(filename: str) -> str:
    base = Path(filename).stem
    base = re.sub(r"[^\w\s.-]", "", base, flags=re.UNICODE)
    base = re.sub(r"[\s_]+", "-", base).strip(".-")
    return (base[:80] if base else "document")


def extract_text(filename: str | None, data: bytes) -> str:
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("File is too large (maximum 8 MB).")
    if not filename:
        raise ValueError("Missing file name.")
    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".text", ".md", ".markdown"):
        text = data.decode("utf-8", errors="replace")
    elif ext == ".docx":
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n\n".join(parts)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    text += "\n" + " | ".join(cells)
    elif ext == ".pdf":
        try:
            reader = PdfReader(BytesIO(data))
        except Exception as e:
            raise ValueError("Could not open this PDF. It may be corrupted, encrypted, or unsupported.") from e
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        text = "\n\n".join(parts)
    else:
        raise ValueError(
            "Unsupported file type. Use .txt, .md, .docx, or .pdf (text-based PDFs only; scanned pages need OCR elsewhere)."
        )

    text = text.strip()
    if not text:
        raise ValueError(
            "No text could be extracted. For PDFs, try a text-based file; scanned images are not read here."
        )
    if len(text) > MAX_INPUT_CHARS:
        text = text[:MAX_INPUT_CHARS]
    return text


def build_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def build_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.replace("\r\n", "\n").split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_bytes(text: str) -> bytes:
    font_path = Path(fpdf.__file__).resolve().parent / "font" / "DejaVuSans.ttf"
    if not font_path.is_file():
        raise RuntimeError("PDF font files missing from fpdf2 package.")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_font("DejaVu", "", str(font_path))
    pdf.set_font("DejaVu", size=11)
    pdf.add_page()
    for line in text.replace("\r\n", "\n").split("\n"):
        if line.strip() == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 6, txt=line)
            pdf.ln(1)
    out = pdf.output(dest="S")
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


def build_file(format_: OutputFormat, text: str) -> tuple[bytes, str]:
    if format_ == "txt":
        return build_txt_bytes(text), "text/plain; charset=utf-8"
    if format_ == "docx":
        return build_docx_bytes(text), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if format_ == "pdf":
        return build_pdf_bytes(text), "application/pdf"
    raise ValueError(f"Unknown output format: {format_}")


def download_part(format_: OutputFormat, text: str, stem: str) -> dict[str, str]:
    raw, media = build_file(format_, text)
    ext = {"txt": "txt", "docx": "docx", "pdf": "pdf"}[format_]
    filename = f"{stem}_plain-language.{ext}"
    return {
        "filename": filename,
        "media_type": media,
        "content_base64": base64.b64encode(raw).decode("ascii"),
    }
